"""
PDF and DOCX extraction utilities.
Handles text extraction, table detection, and page counting.
"""

import io
from typing import Tuple, List, Dict, Any


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, List[Dict], int]:
    """
    Extract text, tables, and page count from a PDF.
    Returns (text, tables, page_count).
    """
    text_parts = []
    tables = []
    page_count = 0

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

                # Extract tables
                try:
                    page_tables = page.extract_tables()
                    for t_idx, table in enumerate(page_tables):
                        if table:
                            table_data = {
                                "page": page_num,
                                "table_index": t_idx,
                                "rows": table,
                                "summary": f"Table on page {page_num} ({len(table)} rows)"
                            }
                            tables.append(table_data)
                except Exception:
                    pass   # Tables are optional

    except ImportError:
        # Fallback: try PyPDF2 if pdfplumber not available
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages, 1):
                t = page.extract_text()
                if t:
                    text_parts.append(f"[Page {i}]\n{t}")
        except Exception as e:
            raise RuntimeError(f"Could not extract PDF text: {e}")
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")

    full_text = "\n\n".join(text_parts)

    # If no text extracted, PDF may be scanned — return empty for OCR fallback
    if not full_text.strip():
        full_text = "[PDF appears to be scanned. OCR processing recommended.]"

    return full_text, tables, page_count


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file preserving structure.
    """
    try:
        from docx import Document
        import io as _io

        doc = Document(_io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                # Detect headings
                if para.style.name.startswith("Heading"):
                    level = para.style.name.split()[-1] if para.style.name.split()[-1].isdigit() else "1"
                    parts.append(f"\n{'#' * int(level)} {para.text}\n")
                else:
                    parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                parts.append("\n[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]\n")

        return "\n".join(parts)

    except ImportError:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def extract_tables_from_pdf(file_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Extract only tables from a PDF file.
    Returns list of table dictionaries with page info and row data.
    """
    _, tables, _ = extract_text_from_pdf(file_bytes)
    return tables


def chunk_text(text: str, chunk_size: int = 3000, overlap: int = 200) -> List[str]:
    """
    Split a long text into overlapping chunks for processing.
    Tries to split at paragraph boundaries.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = []
    current_size = 0

    for para in paragraphs:
        para_size = len(para)

        if current_size + para_size > chunk_size and current_chunk:
            # Save current chunk
            chunks.append("\n\n".join(current_chunk))

            # Keep last paragraph for overlap context
            overlap_text = current_chunk[-1] if current_chunk else ""
            current_chunk = [overlap_text] if len(overlap_text) < overlap else []
            current_size = len(overlap_text) if current_chunk else 0

        current_chunk.append(para)
        current_size += para_size

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks
